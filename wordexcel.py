import openpyxl
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


wb = openpyxl.load_workbook('D:\\example.xlsx') #读取表格
s = wb.active  #取得活动表
s_rows = s.max_row  #取得表格最大行数

d = Document()
phh = d.add_paragraph()
ph = phh.add_run('运通、国发2020年中期员工绩效考核表')  #设置标题
p = d.add_paragraph()  #设置副标题行
bmid = p.add_run('部门：')
bm = p.add_run('办公室          ')
gwid = p.add_run('岗位：')
gw = p.add_run('主任          ')
xmid = p.add_run('姓名：')
xm = p.add_run('何庆')


phh.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  #标题居中
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER



t = d.add_table(s_rows-2, 8, style='Table Grid')  #根据行数增加表格
t.alignment = WD_TABLE_ALIGNMENT.CENTER  #表格居中

for i in range(8):   # 标题行
    t.cell(0, i).text = s.cell(3, i+1).value



for x in range(3, 15):
    for y in range(1, 4):
        t.cell(x-3, y-1).text = str(s.cell(x, y).value)  # 读取EXCEL表中前三列数据，复制到WORD表中


t.cell(0, 0).width = Cm(2)  #设置列宽
t.cell(0, 1).width = Cm(8)
t.cell(0, 2).width = Cm(0.6)
t.cell(0, 3).width = Cm(5)
t.cell(0, 4).width = Cm(3)
t.cell(0, 5).width = Cm(0.6)
t.cell(0, 6).width = Cm(1)
t.cell(0, 7).width = Cm(0.6)

for cell in t.rows[0].cells:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #垂直居中
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER    #左右居中


tb = t.rows[-1]
tb.cells[0].text = '沟通记录'
tb.cells[1].text = '考核者评价意见( 签名/日期）'
tb.cells[4].text = '被考核者对评价结果的意见：（ 签名/日期）'
tb.cells[1].merge(tb.cells[3])  #合并单元格
tb.cells[4].merge(tb.cells[7])
tb.height = Cm(2)  #设置行高


tb2 = t.rows[-2]
tb2.cells[0].merge(tb2.cells[4])
tb2.cells[0].text = '合计'




s = d.sections[0]  #取得章节
s.orientation = WD_ORIENTATION.LANDSCAPE  #页面设置横向
s.page_height = Cm(21)  #设置纸张高度
s.page_width = Cm(29)   #设置纸张宽度（上三项必须同时设置，否则不生效）

s.left_margin = Cm(1)  #设置页边距  左边距
s.right_margin = Cm(1)
s.top_margin = Cm(1)
s.bottom_margin = Cm(1)

wb.save('d:\\1e.xlsx')
d.save('d:\\2e.docx')
